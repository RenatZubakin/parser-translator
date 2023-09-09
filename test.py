import traceback
import requests
from bs4 import BeautifulSoup
import re
from urllib.parse import urljoin
from googletrans import Translator
import pandas as pd
from docx import Document


def get_text_from_element(tag):
    try:
        if tag.name == 'div':
            return tag.string.strip() if tag.string else ''
        return re.sub(r'\s+', ' ', tag.get_text().strip())
    except requests.exceptions.RequestException as e:
        print(f"Error: {e}")
        return ''


def get_site_urls(url):
    if url in visited_urls or main_url not in url or url.endswith('.jpg') or url.endswith('.png'):
        return []
    visited_urls.add(url)
    try:
        response = requests.get(url)
        page_content = response.content
    except requests.exceptions.RequestException:
        return []
    soup = BeautifulSoup(page_content, 'lxml')
    links = soup.find_all('a')
    all_links = []
    for link in links:
        href = link.get('href')
        if href and not href.startswith('#') and href not in visited_urls:
            absolute_url = urljoin(url, href)
            all_links.append(absolute_url)
            all_links += get_site_urls(absolute_url)
    return all_links


def translate_sens(text):
    translator = Translator()
    translated_text = translator.translate(text, src='ru', dest='en')
    return translated_text.text


main_url = r'https://new.66bit.ru/'
visited_urls = set()
all_pages = get_site_urls(main_url)

tags_to_extract = ["div", "h1", "h2", "h3", "h4", "h5", "h6", "p", "span", "pre", "li", "dt", "dd", "th", "td", "label",
                   "input", "button", "textarea", "legend"]

doc = Document()

for page in visited_urls:
    try:
        res = requests.get(page)
        page_cont = res.content
        soup = BeautifulSoup(page_cont, 'lxml')
        main = soup.find('main')
        df = pd.DataFrame(columns=['Русский', 'Английский'])
        page_tags = main.find_all()
        page_sentences = []
        for tag in page_tags:
            if tag.name in tags_to_extract:
                tag_txt = get_text_from_element(tag)
                if tag_txt not in page_sentences:
                    page_sentences.append(tag_txt)
        for s in page_sentences:
            if s:
                eng_txt = translate_sens(s)
                df = pd.concat([df,pd.DataFrame({'Русский':[s],'Английский':[eng_txt]})], ignore_index=True)
        doc.add_heading(f'Страница {page}', level=1)
        rows, cols = df.shape
        table = doc.add_table(rows=rows + 1, cols=cols)
        table.autofit = True
        for i, col in enumerate(df.columns):
            table.cell(0, i).text = col
        for r in range(rows):
            for c in range(cols):
                table.cell(r + 1, c).text = str(df.iloc[r, c])
    except Exception as e:
        trace = traceback.format_exc()
        print(trace)
        print(str(e))

doc.save('все_таблицы1.docx')

