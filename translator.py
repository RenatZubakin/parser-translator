import traceback
import requests
from bs4 import BeautifulSoup
import re
from urllib.parse import urljoin
from googletrans import Translator
import pandas as pd
from docx import Document


def get_text_from_element(tag):
    try:
        if tag.name == 'div':
            return tag.string.strip() if tag.string else ''
        return re.sub(r'\s+', ' ', tag.get_text().strip())
    except requests.exceptions.RequestException as e:
        print(f"Error: {e}")
        return ''


def get_site_urls(url):
    global visited_urls
    print('обработаем', url)
    print(url in visited_urls)
    if url in visited_urls or 'new.66bit.ru' not in url or url.endswith('.jpg'):
        return []
    visited_urls.add(url)
    try:
        response = requests.get(url)
        page_content = response.content
    except requests.exceptions.RequestException:
        return []
    soup = BeautifulSoup(page_content, 'lxml')
    links = soup.find_all('a')
    all_links = []
    for link in links:
        href = link.get('href')
        if href and not href.startswith('#') and href not in visited_urls:
            absolute_url = urljoin(url, href)
            all_links.append(absolute_url)
            all_links += get_site_urls(absolute_url)
    return all_links


def translate_sens(text):
    translator = Translator()
    translated_text = translator.translate(text, src='ru', dest='en')
    return translated_text.text


main_url = r'https://new.66bit.ru/'
#visited_urls = set()
#all_pages = get_site_urls(main_url)

tags_to_extract = ["div", "h1", "h2", "h3", "h4", "h5", "h6", "p", "span", "pre", "li", "dt", "dd", "th", "td", "label",
                   "input", "button", "textarea", "legend"]

doc = Document()
visited_urls=['https://new.66bit.ru/']
for page in visited_urls:
    try:
        res = requests.get(page)
        page_cont = res.content
        soup = BeautifulSoup(page_cont, 'lxml')
        main = soup.find('main')
        page_sentences = set()
        df = pd.DataFrame(columns=['Русский', 'Английский'])
        for tag in main.find_all():
            if tag.name in tags_to_extract:
                tag_txt = get_text_from_element(tag)
                if tag_txt and tag_txt not in page_sentences:
                    print(str(tag_txt))
                    print()
                    eng_txt = translate_sens(tag_txt)
                    df = pd.concat([df,pd.DataFrame({'Русский':[tag_txt],'Английский':[eng_txt]})], ignore_index=True)

        doc.add_heading(f'Страница {page}', level=1)
        rows, cols = df.shape
        table = doc.add_table(rows=rows + 1, cols=cols)
        table.autofit = True
        for i, col in enumerate(df.columns):
            table.cell(0, i).text = col
        for r in range(rows):
            for c in range(cols):
                table.cell(r + 1, c).text = str(df.iloc[r, c])
    except Exception as e:
        trace = traceback.format_exc()
        print(trace)

doc.save('все_таблицы1.docx')

# for tag in main.find_all():
#   if tag.name in tags_to_extract:
#      tag_text = get_text_from_element(tag)
#     print(tag_text)
